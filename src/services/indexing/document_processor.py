"""
Procesadores de documentos para extracción de texto.

Soporta múltiples formatos: PDF, TXT, MD, DOCX
"""

from abc import ABC, abstractmethod
from pathlib import Path
from typing import List, Dict, Any, Optional
import mimetypes

from src.shared.logging_utils import get_logger

logger = get_logger(__name__)


class DocumentProcessor(ABC):
    """Clase base abstracta para procesadores de documentos."""
    
    @abstractmethod
    def extract_text(self, file_path: Path) -> str:
        """Extrae el texto del documento."""
        pass
    
    @abstractmethod
    def extract_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extrae metadatos del documento."""
        pass
    
    def process(self, file_path: Path) -> Dict[str, Any]:
        """
        Procesa un documento completo.
        
        Returns:
            Dict con text y metadata
        """
        text = self.extract_text(file_path)
        metadata = self.extract_metadata(file_path)
        
        return {
            "text": text,
            "metadata": metadata,
        }


class PDFProcessor(DocumentProcessor):
    """Procesador de archivos PDF."""
    
    def extract_text(self, file_path: Path) -> str:
        """
        Extrae texto de un PDF.
        
        Intenta extraer el texto directamente del PDF.
        """
        try:
            from pypdf import PdfReader
            
            reader = PdfReader(str(file_path))
            text_parts = []
            
            for page_num, page in enumerate(reader.pages, 1):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        # Agregar marcador de página
                        text_parts.append(f"\n\n--- Página {page_num} ---\n\n")
                        text_parts.append(page_text)
                except Exception as e:
                    logger.warning(f"Error extrayendo página {page_num}: {e}")
                    continue
            
            full_text = "".join(text_parts)
            
            if not full_text.strip():
                logger.warning(f"PDF sin texto extraíble: {file_path}")
                return ""
            
            logger.info(f"Extraídas {len(reader.pages)} páginas de {file_path.name}")
            
            return full_text
            
        except ImportError:
            logger.error("pypdf no instalado. Instalar con: pip install pypdf")
            raise
        except Exception as e:
            logger.error(f"Error procesando PDF {file_path}: {e}")
            raise
    
    def extract_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extrae metadatos del PDF."""
        try:
            from pypdf import PdfReader
            
            reader = PdfReader(str(file_path))
            metadata = {
                "pages": len(reader.pages),
                "format": "PDF",
            }
            
            # Metadatos del PDF si existen
            if reader.metadata:
                meta = reader.metadata
                if meta.title:
                    metadata["title"] = meta.title
                if meta.author:
                    metadata["author"] = meta.author
                if meta.subject:
                    metadata["subject"] = meta.subject
                if meta.creator:
                    metadata["creator"] = meta.creator
            
            return metadata
            
        except Exception as e:
            logger.warning(f"Error extrayendo metadatos PDF: {e}")
            return {"pages": 0, "format": "PDF"}


class TextProcessor(DocumentProcessor):
    """Procesador de archivos de texto plano (TXT, MD)."""
    
    ENCODINGS = ["utf-8", "latin-1", "cp1252", "iso-8859-1"]
    
    def extract_text(self, file_path: Path) -> str:
        """Extrae texto de archivo de texto plano."""
        text = None
        
        # Intentar múltiples encodings
        for encoding in self.ENCODINGS:
            try:
                with open(file_path, "r", encoding=encoding) as f:
                    text = f.read()
                logger.debug(f"Archivo leído con encoding {encoding}")
                break
            except UnicodeDecodeError:
                continue
            except Exception as e:
                logger.error(f"Error leyendo archivo {file_path}: {e}")
                raise
        
        if text is None:
            raise ValueError(f"No se pudo leer el archivo con ningún encoding conocido")
        
        logger.info(f"Archivo de texto leído: {file_path.name} ({len(text)} caracteres)")
        
        return text
    
    def extract_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extrae metadatos básicos del archivo de texto."""
        stat = file_path.stat()
        
        extension = file_path.suffix.lower()
        format_name = "Markdown" if extension == ".md" else "Texto plano"
        
        return {
            "format": format_name,
            "size_bytes": stat.st_size,
            "extension": extension,
        }


class DOCXProcessor(DocumentProcessor):
    """Procesador de archivos DOCX."""
    
    def extract_text(self, file_path: Path) -> str:
        """Extrae texto de un archivo DOCX."""
        try:
            from docx import Document
            
            doc = Document(str(file_path))
            text_parts = []
            
            # Extraer párrafos
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Extraer tablas
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)
            
            full_text = "\n\n".join(text_parts)
            
            logger.info(
                f"DOCX procesado: {file_path.name} "
                f"({len(doc.paragraphs)} párrafos, {len(doc.tables)} tablas)"
            )
            
            return full_text
            
        except ImportError:
            logger.error("python-docx no instalado. Instalar con: pip install python-docx")
            raise
        except Exception as e:
            logger.error(f"Error procesando DOCX {file_path}: {e}")
            raise
    
    def extract_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extrae metadatos del DOCX."""
        try:
            from docx import Document
            
            doc = Document(str(file_path))
            
            metadata = {
                "format": "DOCX",
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables),
            }
            
            # Core properties si existen
            core_props = doc.core_properties
            if core_props.title:
                metadata["title"] = core_props.title
            if core_props.author:
                metadata["author"] = core_props.author
            if core_props.subject:
                metadata["subject"] = core_props.subject
            if core_props.created:
                metadata["created"] = core_props.created.isoformat()
            if core_props.modified:
                metadata["modified"] = core_props.modified.isoformat()
            
            return metadata
            
        except Exception as e:
            logger.warning(f"Error extrayendo metadatos DOCX: {e}")
            return {"format": "DOCX"}


def get_processor(file_path: Path, mime_type: Optional[str] = None) -> DocumentProcessor:
    """
    Obtiene el procesador adecuado para un archivo.
    
    Args:
        file_path: Ruta al archivo
        mime_type: Tipo MIME (opcional, se puede detectar)
        
    Returns:
        Instancia del procesador apropiado
        
    Raises:
        ValueError: Si el formato no está soportado
    """
    if mime_type is None:
        mime_type, _ = mimetypes.guess_type(str(file_path))
    
    extension = file_path.suffix.lower()
    
    # Mapeo de tipos MIME y extensiones a procesadores
    if mime_type == "application/pdf" or extension == ".pdf":
        return PDFProcessor()
    
    elif extension in [".txt", ".md"] or mime_type in ["text/plain", "text/markdown"]:
        return TextProcessor()
    
    elif (
        mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        or extension == ".docx"
    ):
        return DOCXProcessor()
    
    else:
        raise ValueError(
            f"Formato no soportado: {mime_type or extension}. "
            f"Soportados: PDF, TXT, MD, DOCX"
        )


def process_document(file_path: Path, mime_type: Optional[str] = None) -> Dict[str, Any]:
    """
    Procesa un documento y extrae su texto y metadatos.
    
    Args:
        file_path: Ruta al archivo
        mime_type: Tipo MIME (opcional)
        
    Returns:
        Dict con 'text' y 'metadata'
        
    Raises:
        ValueError: Si el formato no está soportado
        Exception: Si hay error procesando el documento
    """
    processor = get_processor(file_path, mime_type)
    return processor.process(file_path)
